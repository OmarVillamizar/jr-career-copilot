"""
Conversor Markdown → DOCX con formato Harvard tradicional optimizado para ATS.

Convenciones de Markdown soportadas:
- `# Nombre`                    → Nombre centrado grande (serif)
- Párrafo inmediato después de H1 → contacto centrado
- `## SECCIÓN`                  → Header en MAYÚSCULAS con línea horizontal completa
- `### Título | Lado derecho`   → Encabezado de entrada con texto a izquierda y derecha
                                  (la barra `|` separa columnas; lado derecho va en negrita)
- Línea `*texto*` o `_texto_`   → línea en cursiva (subtítulos, fechas)
- Línea con `**Clave:** valor`  → línea de habilidad por categoría (sin viñeta)
- `- item` / `* item`           → bullet point (solo para descripciones/logros)
- `---`                         → separador (espacio extra)

ATS-safe: usa párrafos con tab stops (NO tablas), una sola columna lógica,
fuentes estándar (Cambria/Calibri), sin imágenes ni cuadros de texto.
"""

import os
import re
from typing import Optional


# ─────────────────────────────────────────────────────────────────────────────
# Configuración tipográfica
# ─────────────────────────────────────────────────────────────────────────────
FONT_HEADING_SERIF = "Cambria"     # Nombre principal (serif, look Harvard)
FONT_BODY = "Cambria"              # Cuerpo en serif para look clásico Harvard
FONT_SECTION = "Calibri"           # Headers de sección en sans (contraste)

SIZE_NAME = 22                     # pt — nombre
SIZE_CONTACT = 10                  # pt — línea de contacto
SIZE_SECTION = 11                  # pt — headers de sección (MAYÚSCULAS)
SIZE_ENTRY = 11                    # pt — títulos de entrada (puesto, escuela)
SIZE_BODY = 10.5                   # pt — texto general
SIZE_SMALL = 10                    # pt — fechas, ubicaciones

COLOR_BLACK = "000000"
COLOR_RULE = "000000"              # línea bajo headers


def generate_docx(md_content: str, output_path: str) -> Optional[str]:
    """
    Convierte contenido Markdown a un archivo .docx con estilo Harvard ATS-friendly.

    Args:
        md_content: Contenido del CV en formato Markdown.
        output_path: Ruta donde guardar el archivo .docx.

    Returns:
        Ruta del archivo generado, o None si falló.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    except ImportError as exc:
        print(f"[WARN] No se pudo generar DOCX — dependencias faltantes: {exc}")
        print("       Instala con: pip install python-docx")
        return None

    try:
        doc = Document()
        _setup_document(doc)

        blocks = _parse_markdown(md_content)
        _render_blocks(doc, blocks)

        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        doc.save(output_path)
        return output_path

    except Exception as exc:
        print(f"[WARN] Error al generar DOCX: {exc}")
        import traceback
        traceback.print_exc()
        return None


# ─────────────────────────────────────────────────────────────────────────────
# Configuración del documento (márgenes + estilo Normal + tab stops globales)
# ─────────────────────────────────────────────────────────────────────────────
def _setup_document(doc) -> None:
    from docx.shared import Pt, Cm

    # Márgenes ajustados para look Harvard. Mantengo márgenes laterales
    # generosos para preservar el look elegante; los verticales se reducen
    # ligeramente para que un CV típico de 1 puesto quepa en una página.
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Estilo Normal: body
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(SIZE_BODY)
    style.paragraph_format.space_after = Pt(1)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.05


# ─────────────────────────────────────────────────────────────────────────────
# Parser de Markdown a bloques estructurados
# ─────────────────────────────────────────────────────────────────────────────
def _parse_markdown(md: str) -> list:
    """
    Convierte markdown a una lista de bloques tipados.
    Cada bloque es un dict con 'type' y campos relevantes.
    """
    lines = md.split("\n")
    blocks = []
    i = 0

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Línea vacía → ignorar (los bloques manejan su propio espaciado)
        if not line.strip():
            i += 1
            continue

        # H1: nombre
        m = re.match(r"^#\s+(.+)$", line)
        if m:
            name = _strip_md_inline(m.group(1))
            # Recolectar líneas siguientes no vacías y no-encabezado → contacto
            contact_lines = []
            j = i + 1
            while j < len(lines):
                nxt = lines[j].strip()
                if not nxt:
                    j += 1
                    if contact_lines:
                        break
                    continue
                if nxt.startswith("#") or nxt.startswith("---"):
                    break
                contact_lines.append(_strip_md_inline(nxt))
                j += 1
            blocks.append({"type": "name", "text": name})
            if contact_lines:
                blocks.append({"type": "contact", "text": " · ".join(contact_lines)})
            i = j
            continue

        # H2: sección
        m = re.match(r"^##\s+(.+)$", line)
        if m:
            title = _strip_md_inline(m.group(1)).upper()
            blocks.append({"type": "section", "text": title})
            i += 1
            continue

        # H3: entrada (puede tener "|" para layout izquierda/derecha)
        m = re.match(r"^###\s+(.+)$", line)
        if m:
            content = _strip_md_inline(m.group(1))
            if "|" in content:
                left, right = [s.strip() for s in content.split("|", 1)]
            else:
                left, right = content, ""
            blocks.append({"type": "entry_title", "left": left, "right": right})
            i += 1
            continue

        # Separador horizontal
        if re.match(r"^---+\s*$", line):
            blocks.append({"type": "spacer"})
            i += 1
            continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                item_text = re.sub(r"^\s*[-*]\s+", "", lines[i].rstrip())
                items.append(item_text)
                i += 1
            blocks.append({"type": "bullets", "items": items})
            continue

        # Párrafo: ¿es una línea con "|"? → entry_meta (subtítulo con derecha)
        # ¿es una línea de tipo "**X:** y"? → skill_line
        # Si no → paragraph normal
        if "|" in line and not line.startswith(">"):
            left, right = [s.strip() for s in line.split("|", 1)]
            # Recolectar líneas adicionales del párrafo
            blocks.append({"type": "entry_meta", "left": left, "right": right})
            i += 1
            continue

        # Línea con patrón "**Categoría:** valores" → línea de habilidad
        if re.match(r"^\*\*[^*]+:\*\*", line):
            blocks.append({"type": "skill_line", "text": line})
            i += 1
            continue

        # Párrafo normal: acumular líneas hasta blanco/encabezado
        para_lines = [line]
        j = i + 1
        while j < len(lines):
            nxt = lines[j]
            if not nxt.strip():
                break
            if re.match(r"^(#|---|\s*[-*]\s)", nxt):
                break
            if "|" in nxt:
                break
            para_lines.append(nxt.rstrip())
            j += 1
        blocks.append({"type": "paragraph", "text": " ".join(para_lines)})
        i = j

    return blocks


def _strip_md_inline(text: str) -> str:
    """Quita marcas inline simples (**, *, _, `) dejando el texto."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


def _parse_inline_runs(text: str) -> list:
    """
    Parsea **bold**, *italic*, _italic_ en una línea y retorna
    una lista de (texto, bold, italic).
    """
    runs = []
    # Tokenizar por marcas; uso un enfoque simple con regex iterativa
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|_[^_]+_)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], False, False))
        token = m.group(0)
        if token.startswith("**"):
            runs.append((token[2:-2], True, False))
        elif token.startswith("*"):
            runs.append((token[1:-1], False, True))
        elif token.startswith("_"):
            runs.append((token[1:-1], False, True))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False, False))
    return runs or [(text, False, False)]


# ─────────────────────────────────────────────────────────────────────────────
# Renderizado de bloques a Word
# ─────────────────────────────────────────────────────────────────────────────
def _render_blocks(doc, blocks: list) -> None:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for block in blocks:
        btype = block["type"]

        if btype == "name":
            _render_name(doc, block["text"])
        elif btype == "contact":
            _render_contact(doc, block["text"])
        elif btype == "section":
            _render_section(doc, block["text"])
        elif btype == "entry_title":
            _render_two_column_line(
                doc, block["left"], block["right"],
                left_bold=True, right_bold=True,
                size=SIZE_ENTRY, space_before=3,
                font=FONT_SECTION,  # entry titles en sans para destacar
            )
        elif btype == "entry_meta":
            _render_two_column_line(
                doc, block["left"], block["right"],
                left_bold=False, right_bold=False,
                right_italic=True,
                size=SIZE_SMALL, space_before=0,
            )
        elif btype == "skill_line":
            _render_skill_line(doc, block["text"])
        elif btype == "paragraph":
            _render_paragraph(doc, block["text"])
        elif btype == "bullets":
            _render_bullets(doc, block["items"])
        elif btype == "spacer":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)


def _render_name(doc, name: str) -> None:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name)
    run.bold = True
    run.font.name = FONT_HEADING_SERIF
    run.font.size = Pt(SIZE_NAME)


def _render_contact(doc, text: str) -> None:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(SIZE_CONTACT)
    # Línea horizontal debajo del contacto (separador del header)
    _add_bottom_border(p, size=6, color=COLOR_RULE)


def _render_section(doc, title: str) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.name = FONT_SECTION
    run.font.size = Pt(SIZE_SECTION)
    # Línea horizontal debajo del header de sección
    _add_bottom_border(p, size=6, color=COLOR_RULE)


def _render_two_column_line(
    doc, left: str, right: str,
    left_bold: bool = False, right_bold: bool = False,
    right_italic: bool = False,
    size: float = SIZE_BODY, space_before: int = 0,
    font: str = None,
) -> None:
    """
    Renderiza una línea con texto a la izquierda y a la derecha del mismo párrafo,
    usando un tab stop alineado a la derecha. Esto es ATS-safe (un solo párrafo,
    una sola columna lógica) pero visualmente da efecto de dos columnas.
    """
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_TAB_ALIGNMENT

    use_font = font or FONT_BODY

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(1)

    # Tab stop alineado a la derecha en el límite del área de texto.
    # Ancho de página US Letter (8.5") menos márgenes (2cm + 2cm) ≈ 17 cm.
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17), WD_TAB_ALIGNMENT.RIGHT)

    # Renderizar texto izquierdo (puede tener marcas inline)
    for txt, b, it in _parse_inline_runs(left):
        run = p.add_run(txt)
        run.font.name = use_font
        run.font.size = Pt(size)
        run.bold = b or left_bold
        run.italic = it

    if right:
        p.add_run("\t")
        for txt, b, it in _parse_inline_runs(right):
            run = p.add_run(txt)
            run.font.name = use_font
            run.font.size = Pt(size)
            run.bold = b or right_bold
            run.italic = it or right_italic


def _render_skill_line(doc, text: str) -> None:
    """
    Renderiza una línea de habilidades del tipo:
    **Lenguajes:** Java, Python, JavaScript.
    Manteniendo bold en la categoría.
    """
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)

    for txt, bold, italic in _parse_inline_runs(text):
        run = p.add_run(txt)
        run.font.name = FONT_BODY
        run.font.size = Pt(SIZE_BODY)
        run.bold = bold
        run.italic = italic


def _render_paragraph(doc, text: str) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)

    for txt, bold, italic in _parse_inline_runs(text):
        run = p.add_run(txt)
        run.font.name = FONT_BODY
        run.font.size = Pt(SIZE_BODY)
        run.bold = bold
        run.italic = italic


def _render_bullets(doc, items: list) -> None:
    from docx.shared import Pt, Cm

    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)

        # Borrar runs vacíos generados por el estilo y poner los nuestros
        for r in list(p.runs):
            r.text = ""

        for txt, bold, italic in _parse_inline_runs(item):
            run = p.add_run(txt)
            run.font.name = FONT_BODY
            run.font.size = Pt(SIZE_BODY)
            run.bold = bold
            run.italic = italic


# ─────────────────────────────────────────────────────────────────────────────
# Utilidad: borde inferior en un párrafo (línea horizontal estilo Harvard)
# ─────────────────────────────────────────────────────────────────────────────
def _add_bottom_border(paragraph, size: int = 6, color: str = "000000") -> None:
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): str(size),
        qn("w:space"): "1",
        qn("w:color"): color,
    })
    pBdr.append(bottom)
    pPr.append(pBdr)